"""STAGE 5 — RENDER. ResumeDraft -> formatted .docx. Deterministic, no LLM.

Formatting follows the house style in "(C) HANDOFF-resume-docx-format.md":
Calibri body, navy #1F3864 headings/name, grey #444444 sub-text, right-tab dates
at 7.3in. Section order: Profile, Education, Projects, Technical Skills, Experience.

A `Style` (font/spacing scale) lets the fit loop (stages.fillcheck) tighten the
document toward one page without changing any content. Style() defaults reproduce
the spec exactly. The base geometry constants below are the single source of truth
shared with the estimator.

NAME and CONTACT are static identity, not part of the LLM draft. Edit them for a
different person.
"""
from __future__ import annotations
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from schemas import ResumeDraft
from paths import RESUMES_DIR

NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x44, 0x44, 0x44)
RIGHT_TAB = Inches(7.3)

NAME = "TREVOR SITHISAKD"
CONTACT = "sithisakdt@gmail.com  ·  0404 572 597  ·  Sydney, NSW  ·  github.com/TrevorSithisakd"

# --- Base geometry (spec) — shared source of truth with the estimator --------
PAGE_H_PT, PAGE_W_PT = 792.0, 612.0                 # Letter
MARGIN = dict(top=0.45, bottom=0.4, left=0.6, right=0.6)   # inches
# per-element (pt size, space_before, space_after) at Style() = 1.0
NAME_PT, NAME_AFTER = 19, 1
TAGLINE_PT, TAGLINE_AFTER = 10, 1
CONTACT_PT, CONTACT_AFTER = 9, 3
HEADING_PT, HEADING_BEFORE, HEADING_AFTER = 11, 7.5, 3
BODY_PT, BODY_AFTER = 10, 2
SUB_PT, SUB_AFTER = 9, 1
ROLE_PT, ROLE_BEFORE, ROLE_AFTER = 10, 3, 0
SKILL_PT, SKILL_AFTER = 10, 1.5
BULLET_PT, BULLET_AFTER, BULLET_LINE, BULLET_INDENT_IN = 10, 3.0, 1.05, 0.2
MIN_FONT_PT = 8.5                                    # never scale text below this


@dataclass
class Style:
    """Tightening knobs for the fit loop. 1.0 = the spec exactly."""
    font_scale: float = 1.0
    spacing_scale: float = 1.0

    def fs(self, pt: float) -> float:
        return max(MIN_FONT_PT, pt * self.font_scale)

    def sp(self, pt: float) -> float:
        return pt * self.spacing_scale


def _new_doc(st: Style) -> Document:
    doc = Document()
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(st.fs(BODY_PT))
    n.paragraph_format.space_after = Pt(0)
    n.paragraph_format.line_spacing = 1.0
    s = doc.sections[0]
    s.top_margin, s.bottom_margin = Inches(MARGIN["top"]), Inches(MARGIN["bottom"])
    s.left_margin, s.right_margin = Inches(MARGIN["left"]), Inches(MARGIN["right"])
    return doc


def _name_block(doc, tagline, st: Style):
    p = doc.add_paragraph(); r = p.add_run(NAME)
    r.bold = True; r.font.size = Pt(st.fs(NAME_PT)); r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(st.sp(NAME_AFTER))
    if tagline:
        p2 = doc.add_paragraph(); r2 = p2.add_run(tagline)
        r2.bold = True; r2.font.size = Pt(st.fs(TAGLINE_PT)); r2.font.color.rgb = GREY
        p2.paragraph_format.space_after = Pt(st.sp(TAGLINE_AFTER))
    p3 = doc.add_paragraph(); r3 = p3.add_run(CONTACT)
    r3.font.size = Pt(st.fs(CONTACT_PT)); r3.font.color.rgb = GREY
    p3.paragraph_format.space_after = Pt(st.sp(CONTACT_AFTER))


def _heading(doc, text, st: Style):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(st.sp(HEADING_BEFORE))
    p.paragraph_format.space_after = Pt(st.sp(HEADING_AFTER))
    r = p.add_run(text.upper()); r.bold = True
    r.font.size = Pt(st.fs(HEADING_PT)); r.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6")
    b.set(qn("w:space"), "1"); b.set(qn("w:color"), "1F3864")
    pbdr.append(b); pPr.append(pbdr)


def _body(doc, text, st: Style, after=BODY_AFTER):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(st.sp(after))
    r = p.add_run(text); r.font.size = Pt(st.fs(BODY_PT))


def _sub_line(doc, text, st: Style):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(st.sp(SUB_AFTER))
    r = p.add_run(text); r.italic = True; r.font.size = Pt(st.fs(SUB_PT)); r.font.color.rgb = GREY


def _role_line(doc, left_bold, left_rest, right, st: Style):
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
    p.paragraph_format.space_before = Pt(st.sp(ROLE_BEFORE))
    p.paragraph_format.space_after = Pt(st.sp(ROLE_AFTER))
    r = p.add_run(left_bold); r.bold = True; r.font.size = Pt(st.fs(ROLE_PT))
    if left_rest:
        p.add_run(left_rest).font.size = Pt(st.fs(ROLE_PT))
    if right:
        r3 = p.add_run("\t" + right)
        r3.italic = True; r3.font.size = Pt(st.fs(SUB_PT)); r3.font.color.rgb = GREY


def _bullet(doc, text, st: Style):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(st.sp(BULLET_AFTER))
    p.paragraph_format.left_indent = Inches(BULLET_INDENT_IN)
    p.paragraph_format.line_spacing = BULLET_LINE
    r = p.add_run(text); r.font.size = Pt(st.fs(BULLET_PT))


def _skill_line(doc, label, rest, st: Style):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(st.sp(SKILL_AFTER))
    r = p.add_run(label + ": "); r.bold = True; r.font.size = Pt(st.fs(SKILL_PT))
    p.add_run(rest).font.size = Pt(st.fs(SKILL_PT))


def render(draft: ResumeDraft, job_id: int, version, style: Style | None = None) -> Path:
    """ResumeDraft -> styled .docx at data/resumes/<job_id>/v<version>.docx."""
    st = style or Style()
    path = RESUMES_DIR / str(job_id) / f"v{version}.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = _new_doc(st)

    _name_block(doc, draft.tagline, st)
    _heading(doc, "Profile", st)
    _body(doc, draft.profile, st)
    _heading(doc, "Education", st)
    _body(doc, draft.education, st, after=1)

    _heading(doc, "Projects", st)
    for p in draft.projects:
        _role_line(doc, p.title, f":  {p.angle}" if p.angle else "", p.year, st)
        sub = "  ·  ".join(x for x in (p.stack, p.url) if x)
        if sub:
            _sub_line(doc, sub, st)
        for b in p.bullets:
            _bullet(doc, b, st)
    if draft.additional:
        _body(doc, draft.additional, st, after=1)

    _heading(doc, "Technical Skills", st)
    for s in draft.skills:
        _skill_line(doc, s.label, s.content, st)

    if draft.experience:
        _heading(doc, "Experience", st)
        for e in draft.experience:
            _role_line(doc, e.role, f"  |  {e.org}" if e.org else "", e.dates, st)
            for b in e.bullets:
                _bullet(doc, b, st)

    doc.save(str(path))
    return path
